#coding:utf-8

import docx
import glob, os
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt
from docx.oxml.ns import qn


path_now = os.getcwd()
path_captures = path_now + "\captures"


pic_height = 3880000

#valid format
valid_images = [".jpg",".JPG",".png",".PNG"]

#put images' path into list
image_file_name = [fn for fn in os.listdir(path_captures)
              if any(fn.endswith(ext) for ext in valid_images)]
#sort
image_file_name.sort()

image_file_name_noext =[]

for filenames in image_file_name:
    image_file_name_noext.append(filenames.split('.')[0])

print(image_file_name_noext)

image_file_path = []
for fn in image_file_name:
    image_file_path.append(os.path.join(path_now,"captures",fn))



#document = docx.Document()
document = docx.Document('template(dont-touch-it).docx')
#edit margin
section = document.sections[0]
section.left_margin=Cm(1.27)
section.right_margin=Cm(1.27)
section.top_margin=Cm(1.27)
section.bottom_margin=Cm(1.27)

#insert title
p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(u"title")
font = run.font
font.bold = True
font.name= 'New Times Roman'   #English font
font.size=Pt(12)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')  #Chinese font

p.bold = True




#draw table
if len(image_file_path)%2 == 0:
    tbl = document.add_table(rows=len(image_file_path), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
else:
    tbl = document.add_table(rows=len(image_file_path)+1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
	

#start to insert images and numbers
if len(image_file_path)%2 == 0:
    #image even
    #first column images
    for r in range(0,len(image_file_path)-1,2):
         cell = tbl.cell(r, 0)
         pic = image_file_path[r]
         paragraph = cell.paragraphs[0]
         run = paragraph.add_run()
         run.add_picture(pic, height = pic_height)
         cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #first column numbers
    for r in range(1,len(image_file_path),2):
         cell = tbl.cell(r, 0)
         paragraph = cell.paragraphs[0]
         run = paragraph.add_run("%s" %(image_file_name_noext[r-1]))
         cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
         
    #second colunm images
    for r in range(0,len(image_file_path)-1,2):
         cell = tbl.cell(r, 1)
         pic = image_file_path[r+1]
         paragraph = cell.paragraphs[0]
         run = paragraph.add_run()
         run.add_picture(pic, height = pic_height)
         cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #second column numbers
    for r in range(1,len(image_file_path),2):
         cell = tbl.cell(r, 1)
         paragraph = cell.paragraphs[0]
         run = paragraph.add_run("%s" %(image_file_name_noext[r]))
         cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

else:
    #image odd
    #first column images, without last one
    for r in range(0,len(image_file_path)-2,2):
         cell = tbl.cell(r, 0)
         pic = image_file_path[r]
         paragraph = cell.paragraphs[0]
         run = paragraph.add_run()
         run.add_picture(pic, height = pic_height)
         cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
         
    #first column numbers, without last one
    for r in range(1,len(image_file_path)-1,2):
         cell = tbl.cell(r, 0)
         paragraph = cell.paragraphs[0]
         run = paragraph.add_run("%s" %(image_file_name_noext[r-1]))
         cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

         
    #second column images     
    for r in range(0,len(image_file_path)-1,2):
         cell = tbl.cell(r, 1)
         pic = image_file_path[r+1]
         paragraph = cell.paragraphs[0]
         run = paragraph.add_run()
         run.add_picture(pic, height = pic_height)
         cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #second column numbers
    for r in range(1,len(image_file_path)-1,2):
         cell = tbl.cell(r, 1)
         paragraph = cell.paragraphs[0]
         run = paragraph.add_run("%s" %(image_file_name_noext[r]))
         cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

         
    #for the last cell
    cell = tbl.cell(len(image_file_path), 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("%s" %(len(image_file_path)))
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #for the last cell
    cell = tbl.cell(len(image_file_path)-1, 0)
    pic = image_file_path[len(image_file_path)-1]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(pic, height = pic_height)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER



document.save("output.docx")
